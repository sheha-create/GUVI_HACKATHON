"""
File parser utility for extracting text from various document formats.
Supports PDF, DOCX, and image files (JPG, PNG).
"""

import fitz  # PyMuPDF
import pdfplumber
from docx import Document
from PIL import Image
import pytesseract
import os
from typing import Tuple


def extract_text_from_pdf(file_path: str) -> str:
    """
    Extract text from PDF file using PyMuPDF and pdfplumber.
    
    Args:
        file_path: Path to the PDF file
        
    Returns:
        Extracted text from the PDF
    """
    text = ""
    
    try:
        # Try pdfplumber first (better for structured data)
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                text += page.extract_text() or ""
                text += "\n"
    except Exception as e:
        print(f"pdfplumber extraction failed: {e}, trying PyMuPDF...")
        # Fallback to PyMuPDF
        try:
            pdf_doc = fitz.open(file_path)
            for page in pdf_doc:
                text += page.get_text()
                text += "\n"
            pdf_doc.close()
        except Exception as e2:
            raise Exception(f"Failed to extract text from PDF: {str(e2)}")
    
    return text.strip()


def extract_text_from_docx(file_path: str) -> str:
    """
    Extract text from DOCX file.
    
    Args:
        file_path: Path to the DOCX file
        
    Returns:
        Extracted text from the DOCX
    """
    try:
        doc = Document(file_path)
        text = ""
        
        # Extract text from paragraphs
        for para in doc.paragraphs:
            text += para.text + "\n"
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
                text += "\n"
        
        return text.strip()
    except Exception as e:
        raise Exception(f"Failed to extract text from DOCX: {str(e)}")


def extract_text_from_image(file_path: str) -> str:
    """
    Extract text from image file using OCR (Tesseract).
    Falls back to empty string if Tesseract is not installed.
    
    Args:
        file_path: Path to the image file
        
    Returns:
        Extracted text from the image
    """
    try:
        # Try to use Tesseract OCR if available
        image = Image.open(file_path)
        text = pytesseract.image_to_string(image)
        return text.strip()
    except Exception as e:
        print(f"OCR extraction failed (Tesseract not installed): {e}")
        # Return a note that OCR is not available
        return "[OCR not available - Tesseract not installed. Please install Tesseract for image text extraction.]"


def parse_document(file_path: str, file_type: str) -> Tuple[str, str]:
    """
    Parse a document based on its file type and extract text.
    
    Args:
        file_path: Path to the document file
        file_type: Type of file (pdf, docx, jpg, png)
        
    Returns:
        Tuple of (extracted_text, document_type_hint)
    """
    file_type = file_type.lower().strip(".")
    
    if file_type in ["pdf"]:
        return extract_text_from_pdf(file_path), "PDF Document"
    
    elif file_type in ["docx", "doc"]:
        return extract_text_from_docx(file_path), "Word Document"
    
    elif file_type in ["jpg", "jpeg", "png", "bmp", "gif"]:
        return extract_text_from_image(file_path), "Image Document"
    
    else:
        raise ValueError(f"Unsupported file type: {file_type}")


def validate_file_type(filename: str) -> bool:
    """
    Validate if the file type is supported.
    
    Args:
        filename: Name of the file
        
    Returns:
        True if file type is supported, False otherwise
    """
    supported_types = {".pdf", ".docx", ".doc", ".jpg", ".jpeg", ".png", ".bmp", ".gif"}
    file_ext = os.path.splitext(filename)[1].lower()
    return file_ext in supported_types
